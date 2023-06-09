import docx
from util import *
import time

name = 'caudo BAHNAR.docx'

st = time.time()

doc = docx.Document('document/' + name)
all_paras = doc.paragraphs

mydoc = docx.Document()
lines = []
for para in all_paras:
    text = para.text
    for original, replace in radioContentMap.items():
        text = text.replace(original, replace)
    for original, replace in contentSpecialMap.items():
        text = text.replace(original, replace)
    for original, replace in contentMap.items():
        text = text.replace(original, replace)
    lines.append(text)
    
lines_lst = []
for line in lines:
    if len(line) == 0:
        continue
    line = line.split('.')
    for l in line:
        lines_lst.append(l)
    
lines = lines_lst
line_with_type = []
for line in lines:
    line_with_type.append((line,check_sentence(line)))

lines = []
for line in line_with_type:
    str_line,typ = line
    if typ == True:
        lines.append(process_vi(str_line) + '.')
    else:
        lines.append(process_bahnar(str_line) + '.')
with open('correct_' + name, 'w', encoding='utf-8') as f:
    f.writelines(lines) 

et = time.time()
elapsed_time = et - st
print(elapsed_time)
